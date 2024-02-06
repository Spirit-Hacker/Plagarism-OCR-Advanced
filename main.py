from fastapi import FastAPI, File, UploadFile, Form ,HTTPException
from fastapi.responses import HTMLResponse
from algorithm import main
from algorithm import fileSimilarity
from PyPDF2 import PdfReader
import pytesseract
import io
from PIL import Image
from docx import Document
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI()


origins = [
    "http://localhost:5173",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
async def home():
    return {"Endpoint is working"}

@app.post("/text_similarity/")
async def test(q: str = Form(...)):
    print("request is welcome test")
    print(q)
    
    if q:
        percent, link, plagiarized_text = main.findSimilarity(q)  # Assuming main.findSimilarity returns plagiarized_text along with percent and link
        percent = round(percent, 2)
    
    print("Output : ", percent, link, plagiarized_text)
    print('/n')
    return {"link": link, "percent": percent, "plagiarized_text": plagiarized_text}


@app.post("/doc_similarity/")
async def filetest(docfile: UploadFile = File(...)):
    value = ""
    print(docfile.filename)
    if docfile.filename.endswith(".txt"):
        value = (await docfile.read()).decode()

    elif docfile.filename.endswith(".docx"):
        document = Document(io.BytesIO(await docfile.read()))
        for para in document.paragraphs:
            value += para.text

    elif docfile.filename.endswith(".pdf"):
        pdfReader = PdfReader(io.BytesIO(await docfile.read()))
        pageObj = pdfReader.pages[0]
        value = pageObj.extract_text()

    percent, link, plagiarised_text = main.findSimilarity(value)
    print("Output : ", percent, link)
    print('/n')
    return {"link": link, "percent": percent, "Plagiarised_text": plagiarised_text}

@app.post("/image_plagiarism_check/")
async def image_plagiarism_check(image: UploadFile = File(...)):
    pytesseract.pytesseract.tesseract_cmd = "C:/Program Files/Tesseract-OCR/tesseract.exe"
    if image:
        # Read the uploaded image file
        contents = await image.read()
        
        # Open the image using PIL
        uploaded_image = Image.open(io.BytesIO(contents))
        
        # Extract text from the image using pytesseract
        extracted_text = pytesseract.image_to_string(uploaded_image)
        
        # Find similarity with main algorithm
        percent, link, plagiarized_text = main.findSimilarity(extracted_text)
        percent = round(percent, 2)
        
        # Return the result
        return {"link": link, "percent": percent, "plagiarised text": plagiarized_text}
    print(plagiarized_text)
    print('/n')
    
    # If no image is provided, return an empty response
    return {}



@app.post("/two_text_comparison/")
async def twofiletest1(q1: str = Form(...), q2: str = Form(...)):
    print("Submiited text for 1st and 2nd")
    print(q1)
    print(q2)
    if q1 != "" and q2 != "":
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(q1, q2)
    result = round(result, 2)
    print("Output : ", result)
    print('/n')
    return {"result": result}

@app.post("/two_file_comparison/")
async def twofilecompare1(docfile1: UploadFile = File(...), docfile2: UploadFile = File(...)):
    value1 = ""
    value2 = ""
    if docfile1.filename.endswith(".txt") and docfile2.filename.endswith(".txt"):
        value1 = (await docfile1.read()).decode()
        value2 = (await docfile2.read()).decode()

    elif docfile1.filename.endswith(".docx") and docfile2.filename.endswith(".docx"):
        document1 = Document(io.BytesIO(await docfile1.read()))
        document2 = Document(io.BytesIO(await docfile2.read()))
        for para in document1.paragraphs:
            value1 += para.text
        for para in document2.paragraphs:
            value2 += para.text
    
    elif docfile1.filename.endswith(".pdf") and docfile2.filename.endswith(".pdf"):
        pdf_reader1 = PdfReader(io.BytesIO(await docfile1.read()))
        for page in pdf_reader1.pages:
            value1 += page.extract_text()

        pdf_reader2 = PdfReader(io.BytesIO(await docfile2.read()))
        for page in pdf_reader2.pages:
            value2 += page.extract_text()

    else:
        raise HTTPException(status_code=400, detail="Unsupported file types. Please upload TXT, DOCX, or PDF files")
    
    result = fileSimilarity.findFileSimilarity(value1, value2)
    print("Output : ", result)
    print('/n')
    return {"result": result}
